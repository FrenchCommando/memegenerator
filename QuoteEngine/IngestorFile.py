from abc import ABC, abstractmethod
import csv
from docx import Document
import os
from typing import List
from QuoteEngine.QuoteModelFile import QuoteModel
from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser


def convert_pdf_to_string(file_path):
    output_string = StringIO()
    with open(file_path, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        for page in PDFPage.create_pages(doc):
            interpreter.process_page(page)
    return output_string.getvalue()


class IngestorInterface(ABC):

    @staticmethod
    @abstractmethod
    def can_ingest(path: str) -> bool:
        pass

    @staticmethod
    @abstractmethod
    def parse(path: str) -> List[QuoteModel]:
        pass


class IngestorCsv(IngestorInterface):
    @staticmethod
    def can_ingest(path: str) -> bool:
        extension = os.path.splitext(path)[1]
        return extension == ".csv"

    @staticmethod
    def parse(path: str) -> List[QuoteModel]:
        with open(path, mode='r') as input_file:
            reader = csv.DictReader(input_file)
            return [QuoteModel(**row) for row in reader]


class IngestorTxt(IngestorInterface):
    @staticmethod
    def can_ingest(path: str) -> bool:
        extension = os.path.splitext(path)[1]
        return extension == ".txt"

    @staticmethod
    def parse(path: str) -> List[QuoteModel]:
        with open(path, mode='r') as input_file:
            reader = input_file.readlines()
            return [QuoteModel(*map(lambda r: r.strip(),
                                    row.encode('ascii', 'ignore').decode().split("-")))
                    for row in reader]


class IngestorDocx(IngestorInterface):
    @staticmethod
    def can_ingest(path: str) -> bool:
        extension = os.path.splitext(path)[1]
        return extension == ".docx"

    @staticmethod
    def parse(path: str) -> List[QuoteModel]:
        document = Document(path)
        return [QuoteModel(*map(lambda r: r.strip().strip('\"'), para.text.split("-")))
                for para in document.paragraphs if para.text]


class IngestorPdf(IngestorInterface):
    @staticmethod
    def can_ingest(path: str) -> bool:
        extension = os.path.splitext(path)[1]
        return extension == ".pdf"

    @staticmethod
    def parse(path: str) -> List[QuoteModel]:
        quote_list = []
        text = convert_pdf_to_string(file_path=path)
        quote_list.extend([QuoteModel(*map(lambda r: r.strip().strip('\"'), line.split("-")))
                           for line in text.split("\n") if line.strip().encode('ascii', 'ignore').decode()])
        return quote_list


class Ingestor:
    extension_mapping = {
        '.csv': IngestorCsv,
        '.docx': IngestorDocx,
        '.pdf': IngestorPdf,
        '.txt': IngestorTxt,
    }

    @staticmethod
    def parse(path: str) -> List[QuoteModel]:
        extension = os.path.splitext(path)[1]
        if extension not in Ingestor.extension_mapping:
            raise ValueError(f"Ingestor: Extension not supported for extension {extension}, file: {path}")
        ingestor = Ingestor.extension_mapping[extension]
        if ingestor.can_ingest(path=path):  # allows to check more than extension type
            try:
                return ingestor.parse(path=path)
            except BaseException as e:
                raise ValueError(f"Ingestor {ingestor} failed for \t{path}\nwith error:\t{e}")
        raise ValueError(f"Ingestor {ingestor} not implemented for {path}")
